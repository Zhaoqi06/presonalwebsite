import streamlit as st
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

if "logged_in" not in st.session_state or not st.session_state.logged_in:
    st.error("请先登录")
    st.switch_page("pages/login.py")

# 隐藏默认导航/水印
st.markdown("""
            <style>
            .css-14xtw13,.css-1v3fvcr{display:none !important;}
            </style>
              """, unsafe_allow_html=True)

if "刘钊齐" == st.session_state["username"]:
    st.title("欢迎进入管理页面")
    st.divider()
    nav = st.sidebar.selectbox("导航栏", ["首页", "数学教学"])
    if nav == "首页":
        st.header("发布通知")

        # 获取当前脚本所在目录
        script_dir = os.path.dirname(os.path.abspath(__file__))
        docx_file_path = os.path.join(script_dir, "..", "document", "notifications.docx")
        docx_file_path = os.path.normpath(docx_file_path)

        with st.form("input_form"):
            title_input = st.text_input("标题")
            text_input = st.text_area("文本")
            time_input = st.date_input("时间")
            submit_btn = st.form_submit_button("提交")

        if submit_btn:
            if not title_input or not text_input or not time_input:
                st.error("请填写完整的信息")
            else:
                try:
                    # 检查文件是否存在，不存在则创建新文档
                    if os.path.exists(docx_file_path):
                        doc = Document(docx_file_path)
                    else:
                        doc = Document()

                    # 添加标题
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(f"标题: {title_input}")
                    title_run.bold = True
                    title_run.font.size = Pt(14)

                    # 添加内容
                    content_para = doc.add_paragraph()
                    content_para.add_run(f"文本: {text_input}")

                    # 添加时间
                    time_para = doc.add_paragraph()
                    time_run = time_para.add_run(f"时间: {time_input}")
                    time_run.font.size = Pt(10)

                    # 添加分隔线
                    separator_para = doc.add_paragraph()
                    separator_run = separator_para.add_run("-" * 50)
                    separator_run.font.size = Pt(8)

                    # 保存文档
                    doc.save(docx_file_path)
                    st.success("发布成功！")
                except Exception as e:
                    st.error(f"发布失败：{str(e)}")

        st.divider()
        st.header("删除通知")

        with st.form("delete_form"):
            title_input = st.text_input("标题")
            time_input = st.date_input("时间")
            submit_btn = st.form_submit_button("提交")

        # 在删除功能部分使用更精确的匹配
        if submit_btn:
            try:
                if not os.path.exists(docx_file_path):
                    st.error("文档不存在")
                    st.stop()

                doc = Document(docx_file_path)

                # 更精确地查找通知块
                paragraphs_text = [p.text for p in doc.paragraphs]
                paragraphs_to_remove = []

                i = 0
                while i < len(paragraphs_text):
                    # 查找标题段落
                    if paragraphs_text[i].startswith("标题:") and title_input in paragraphs_text[i]:
                        # 检查是否在接下来的几段内有对应的时间
                        for j in range(i, min(i + 4, len(paragraphs_text))):
                            if paragraphs_text[j].startswith("时间:") and str(time_input) in paragraphs_text[j]:
                                # 找到匹配的通知块，标记要删除的段落
                                paragraphs_to_remove.extend(range(i, min(i + 4, len(paragraphs_text))))
                                i = min(i + 4, len(paragraphs_text))  # 跳过已处理的通知块
                                break
                        else:
                            i += 1
                    else:
                        i += 1

                # 从后往前删除段落
                for idx in sorted(paragraphs_to_remove, reverse=True):
                    if idx < len(doc.paragraphs):
                        p = doc.paragraphs[idx]._element
                        p.getparent().remove(p)

                # 保存文档
                doc.save(docx_file_path)
                st.success("通知删除成功！")

            except Exception as e:
                st.error(f"删除失败：{str(e)}")
else:
    st.error("非管理员人员不能进入该页面")
