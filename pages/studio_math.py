import streamlit as st
import requests
import os
import base64
import re
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from PIL import Image

# ==================== 初始化配置 ====================
load_dotenv()
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

st.set_page_config(
    page_title="MathAI · Research Studio",
    page_icon="🧪",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ---------- 会话状态初始化 ----------
defaults = {
    'theme': "dark",
    'math_history': [],
    'paper_text': "",
    'file_name': "",
    'current_result': "",
    'favorites': [],
    'math_search_query': ""
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# 清理可能存在的旧格式历史记录（确保结构一致）
if 'math_history' in st.session_state:
    cleaned = []
    for item in st.session_state.math_history:
        if isinstance(item, dict):
            # 确保必要的键存在
            if 'q' not in item:
                item['q'] = item.get('full_q', '无标题')[:80]
            if 'full_q' not in item:
                item['full_q'] = item.get('q', '')
            if 'a' not in item:
                item['a'] = ''
            cleaned.append(item)
    st.session_state.math_history = cleaned

# ==================== 主题系统 (玻璃态工作室风格) ====================
def get_theme_css(theme):
    glass_light = """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:opsz@14..32&family=JetBrains+Mono&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
    .main { background: #f4f7fb; color: #1e293b; }
    [data-testid="stSidebar"] { background: rgba(255,255,255,0.75); backdrop-filter: blur(12px); border-right: 1px solid rgba(0,0,0,0.05); }
    h1 { color: #0f172a; font-weight: 600; letter-spacing: -0.02em; }
    h2, h3 { color: #1e293b; font-weight: 500; }
    .glass-card {
        background: rgba(255,255,255,0.7); backdrop-filter: blur(12px);
        border: 1px solid rgba(255,255,255,0.8); border-radius: 24px;
        padding: 24px; margin-bottom: 20px;
        box-shadow: 0 20px 35px -8px rgba(0,0,0,0.07), 0 0 0 1px rgba(0,0,0,0.02);
    }
    .stButton>button {
        background: #1e293b; color: white; border: none; border-radius: 60px;
        padding: 10px 24px; font-weight: 500; transition: all 0.2s;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }
    .stButton>button:hover { background: #0f172a; transform: translateY(-1px); box-shadow: 0 10px 20px -5px rgba(0,0,0,0.15); }
    .stTextArea textarea, .stTextInput input, .stSelectbox select {
        background: rgba(255,255,255,0.8); border: 1px solid #e2e8f0; border-radius: 16px;
    }
    .chat-bubble {
        background: rgba(255,255,255,0.6); backdrop-filter: blur(4px);
        border-radius: 20px; padding: 18px 22px; margin: 12px 0;
        border: 1px solid rgba(255,255,255,0.9);
    }
    .user-bubble { border-left: 6px solid #2563eb; }
    .ai-bubble { border-left: 6px solid #7c3aed; }
    .code-block { background: #1e293b; color: #e2e8f0; padding: 16px; border-radius: 16px; font-family: 'JetBrains Mono', monospace; }
    </style>
    """
    glass_dark = """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:opsz@14..32&family=JetBrains+Mono&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
    .main { background: #0b1120; color: #e2e8f0; }
    [data-testid="stSidebar"] { background: rgba(15,23,42,0.7); backdrop-filter: blur(16px); border-right: 1px solid #1e293b; }
    h1 { color: #f8fafc; font-weight: 600; letter-spacing: -0.02em; }
    h2, h3 { color: #cbd5e1; font-weight: 500; }
    .glass-card {
        background: rgba(30,41,59,0.5); backdrop-filter: blur(16px);
        border: 1px solid rgba(51,65,85,0.5); border-radius: 24px;
        padding: 24px; margin-bottom: 20px;
        box-shadow: 0 20px 35px -8px rgba(0,0,0,0.4), 0 0 0 1px rgba(255,255,255,0.02);
    }
    .stButton>button {
        background: #3b82f6; color: white; border: none; border-radius: 60px;
        padding: 10px 24px; font-weight: 500; transition: all 0.2s;
        box-shadow: 0 4px 12px rgba(59,130,246,0.2);
    }
    .stButton>button:hover { background: #2563eb; transform: translateY(-1px); box-shadow: 0 10px 20px -5px #1e3a8a; }
    .stTextArea textarea, .stTextInput input, .stSelectbox select {
        background: rgba(15,23,42,0.6); border: 1px solid #334155; border-radius: 16px; color: #f1f5f9;
    }
    .chat-bubble {
        background: rgba(30,41,59,0.5); backdrop-filter: blur(8px);
        border-radius: 20px; padding: 18px 22px; margin: 12px 0;
        border: 1px solid #334155;
    }
    .user-bubble { border-left: 6px solid #3b82f6; }
    .ai-bubble { border-left: 6px solid #8b5cf6; }
    .code-block { background: #0f172a; color: #cbd5e1; padding: 16px; border-radius: 16px; font-family: 'JetBrains Mono', monospace; }
    </style>
    """
    return glass_dark if theme == "dark" else glass_light

st.markdown(get_theme_css(st.session_state.theme), unsafe_allow_html=True)

# ==================== 工具函数 ====================
def extract_text_from_pdf(file):
    try:
        reader = PdfReader(BytesIO(file.read()))
        return "\n\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
    except Exception:
        return ""

def extract_text_from_docx(file):
    try:
        doc = Document(BytesIO(file.read()))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = [" | ".join(c.text for c in row.cells) for t in doc.tables for row in t.rows]
        return "\n\n".join(paras + tables)
    except Exception:
        return ""

def extract_text_from_txt(file):
    return file.read().decode("utf-8")

def extract_text_from_md(file):
    return file.read().decode("utf-8")

def process_uploaded_file(uploaded_file):
    if not uploaded_file:
        return ""
    ext = uploaded_file.name.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(uploaded_file)
    elif ext == "docx":
        return extract_text_from_docx(uploaded_file)
    elif ext == "txt":
        return extract_text_from_txt(uploaded_file)
    elif ext == "md":
        return extract_text_from_md(uploaded_file)
    else:
        st.error(f"不支持格式: .{ext}")
        return ""

def encode_image(img):
    return base64.b64encode(img.read()).decode()

def call_deepseek(system, user, image=None):
    if not DEEPSEEK_API_KEY:
        st.error("API Key 缺失")
        return None
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    messages = [{"role": "system", "content": system}]
    if image:
        messages.append({"role": "user", "content": [
            {"type": "text", "text": user},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image}"}}
        ]})
        model = "deepseek-vl"
    else:
        messages.append({"role": "user", "content": user})
        model = "deepseek-chat"
    payload = {"model": model, "messages": messages, "temperature": 0.7, "max_tokens": 8000}
    try:
        with st.spinner("🧠 思考中..."):
            r = requests.post("https://api.deepseek.com/v1/chat/completions", headers=headers, json=payload)
            return r.json()["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"请求失败: {e}")
        return None

# ==================== 侧边栏 ====================
with st.sidebar:
    st.markdown("# 🧪 MathAI Studio")
    st.caption("研究级 · 数学智能实验室")
    st.divider()
    theme = st.radio("主题", ["🌙 深色", "☀️ 浅色"], horizontal=True, label_visibility="collapsed")
    if "深色" in theme and st.session_state.theme != "dark":
        st.session_state.theme = "dark"
        st.rerun()
    elif "浅色" in theme and st.session_state.theme != "light":
        st.session_state.theme = "light"
        st.rerun()

    st.markdown("### 📌 工作区")
    page = st.radio("导航", ["🧮 数学实验室", "📄 论文审稿室"], label_visibility="collapsed")

    st.divider()
    if DEEPSEEK_API_KEY:
        st.success("🟢 API 在线")
    else:
        st.error("🔴 未连接")

    if st.session_state.math_history:
        st.metric("历史记录", len(st.session_state.math_history))
    if st.session_state.favorites:
        st.metric("收藏", len(st.session_state.favorites))

# ==================== 数学实验室 ====================
def math_lab():
    st.title("🧮 数学推理实验室")
    st.caption("高等数学 · 建模 · 符号计算 · 代码生成")

    with st.container():
        cols = st.columns(4)
        cols[0].markdown("🔬 **代数/微积分**")
        cols[1].markdown("📈 **概率/优化**")
        cols[2].markdown("🖼️ **图像识别**")
        cols[3].markdown("💻 **代码生成**")

    st.divider()

    tab1, tab2, tab3 = st.tabs(["📝 文本输入", "🖼️ 图片上传", "📂 文件上传 (MD/TXT)"])

    with tab1:
        col1, col2 = st.columns([3, 1])
        with col1:
            problem = st.text_area("输入数学问题", height=140,
                                   placeholder="例如: 求解微分方程 y'' + 2y' + y = 0")
        with col2:
            domain = st.selectbox("领域", ["自动识别", "微积分", "线性代数", "概率统计", "优化理论", "数论"])
            mode = st.selectbox("风格", ["标准", "严谨推导", "教学讲解"])
            render_latex = st.checkbox("✨ 渲染 LaTeX", True)

    with tab2:
        img_file = st.file_uploader("上传数学题图片", type=["png", "jpg", "jpeg"], key="math_img")
        img_b64 = None
        if img_file:
            st.image(img_file, use_container_width=True)
            img_b64 = encode_image(img_file)

    with tab3:
        up_file = st.file_uploader("上传 Markdown 或 TXT 文件", type=["md", "txt"], key="math_file")
        file_content = ""
        if up_file:
            file_content = process_uploaded_file(up_file)
            st.text_area("文件内容预览", file_content, height=150)

    c1, c2, c3 = st.columns([2, 1, 1])
    with c1:
        go = st.button("🚀 开始推理", type="primary", use_container_width=True)
    with c2:
        if st.button("🗑️ 清除历史"):
            st.session_state.math_history = []
            st.rerun()
    with c3:
        if st.button("⭐ 清空收藏"):
            st.session_state.favorites = []
            st.rerun()

    if go:
        user_input = ""
        if tab1 and problem:
            user_input = problem
        elif tab2 and img_b64:
            user_input = "请解答图片中的数学问题"
        elif tab3 and file_content:
            user_input = file_content

        if not user_input and not img_b64:
            st.warning("请提供输入")
            return

        sys_prompt = "你是一位顶尖数学研究员，回答专业、清晰。"
        if domain != "自动识别":
            sys_prompt += f" 专注于{domain}领域。"
        if mode == "严谨推导":
            sys_prompt += " 使用公理化证明，所有公式用 LaTeX 格式。"
        elif mode == "教学讲解":
            sys_prompt += " 用生动比喻，逐步讲解，适合教学。"

        result = call_deepseek(sys_prompt, user_input, img_b64)
        if result:
            st.session_state.math_history.append({
                "q": user_input[:80] + ("…" if len(user_input) > 80 else ""),
                "full_q": user_input,
                "a": result,
                "time": len(st.session_state.math_history)
            })
            st.session_state.current_result = result
            st.rerun()

    if st.session_state.current_result:
        st.divider()
        st.markdown("### 📌 最新结果")
        with st.container():
            st.markdown(f'<div class="glass-card">{st.session_state.current_result}</div>', unsafe_allow_html=True)
            code_pattern = r'```(?:python)?\n(.*?)```'
            codes = re.findall(code_pattern, st.session_state.current_result, re.DOTALL)
            if codes:
                with st.expander("💻 查看代码片段"):
                    for i, c in enumerate(codes):
                        st.code(c, language='python')

        col_dl, col_fav = st.columns([1, 1])
        with col_dl:
            st.download_button("📥 导出 Markdown", st.session_state.current_result, "math_result.md")
        with col_fav:
            if st.button("⭐ 收藏此结果"):
                if st.session_state.current_result not in st.session_state.favorites:
                    st.session_state.favorites.append(st.session_state.current_result)
                    st.success("已收藏")

    if st.session_state.math_history:
        st.divider()
        st.markdown("### 📚 推理历史")
        search = st.text_input("🔍 搜索历史", placeholder="关键词", key="math_search")
        filtered = [h for h in reversed(st.session_state.math_history) if search.lower() in h.get('full_q', '').lower()] if search else reversed(st.session_state.math_history)
        for i, h in enumerate(filtered):
            title = h.get('q', '无标题')
            with st.expander(f"{title}", expanded=(i == 0 and not search)):
                st.markdown(f'<div class="chat-bubble user-bubble"><b>❓ 问题</b><br>{h.get("full_q", "")}</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="chat-bubble ai-bubble"><b>🤖 解答</b><br>{h.get("a", "")}</div>', unsafe_allow_html=True)

    if st.session_state.favorites:
        with st.sidebar:
            st.divider()
            st.markdown("### ⭐ 收藏夹")
            for idx, fav in enumerate(st.session_state.favorites):
                with st.expander(f"收藏 #{idx + 1}"):
                    st.markdown(fav[:300] + "…" if len(fav) > 300 else fav)

# ==================== 论文审稿室 ====================
def paper_review():
    st.title("📄 论文审稿室")
    st.caption("专业评审 · 多维度评分 · 修改建议")

    with st.container():
        st.markdown('<div class="glass-card">支持 PDF, DOCX, TXT, MD 格式</div>', unsafe_allow_html=True)

    up_file = st.file_uploader("上传论文", type=["pdf", "docx", "txt", "md"], key="paper_upload")
    if up_file:
        if up_file.name != st.session_state.file_name:
            with st.spinner("解析中..."):
                st.session_state.paper_text = process_uploaded_file(up_file)
                st.session_state.file_name = up_file.name
            if st.session_state.paper_text:
                st.success(f"已加载 {up_file.name} ({len(st.session_state.paper_text)} 字符)")

    if st.session_state.paper_text and st.button("清除当前论文"):
        st.session_state.paper_text = ""
        st.session_state.file_name = ""
        st.rerun()

    content = st.text_area("论文内容", value=st.session_state.paper_text, height=180, placeholder="也可直接粘贴...")

    c1, c2 = st.columns(2)
    with c1:
        title = st.text_input("论文标题 (可选)")
    with c2:
        field = st.selectbox("研究领域", ["数学", "计算机科学", "控制工程", "统计学", "其他"])

    if st.button("📋 生成审稿报告", type="primary"):
        final_text = st.session_state.paper_text or content
        if not final_text:
            st.warning("请提供论文内容")
            return
        sys = f"你是{field}领域资深审稿人。按以下结构输出：\n## 一、总体概览\n## 二、创新性/科学性/写作评价\n## 三、评分表(0-100)\n## 四、发表建议及具体修改意见"
        res = call_deepseek(sys, f"标题:{title}\n\n{final_text}")
        if res:
            st.session_state.current_result = res
            st.markdown("### 📑 审稿报告")
            st.markdown(f'<div class="glass-card">{res}</div>', unsafe_allow_html=True)
            st.download_button("📥 下载报告", res, "review.md")

# ==================== 主入口 ====================
if page == "🧮 数学实验室":
    math_lab()
else:
    paper_review()